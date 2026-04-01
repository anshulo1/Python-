import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx2pdf import convert
from pdf2docx import Converter
from PyPDF2 import PdfMerger
from docx import Document


# PROCESS WINDOW
def show_processing(task_function):

    process_window = tk.Toplevel(root)
    process_window.title("Processing")
    process_window.geometry("350x150")

    label = tk.Label(process_window, text="Processing file...", font=("Arial", 12))
    label.pack(pady=10)

    progress = ttk.Progressbar(process_window, mode="indeterminate", length=250)
    progress.pack(pady=10)
    progress.start(10)

    cancelled = {"value": False}

    def cancel():
        cancelled["value"] = True
        process_window.destroy()

    cancel_btn = tk.Button(process_window, text="Cancel", command=cancel)
    cancel_btn.pack(pady=10)

    root.update()

    if not cancelled["value"]:
        task_function()

    progress.stop()
    process_window.destroy()


# WORD → PDF
def word_to_pdf():

    file = filedialog.askopenfilename(filetypes=[("Word File", "*.docx")])

    if file:

        save = filedialog.asksaveasfilename(defaultextension=".pdf")

        if save:

            def task():
                convert(file)

            show_processing(task)

            messagebox.showinfo("Success", "Word converted to PDF")


# PDF → WORD
def pdf_to_word():

    file = filedialog.askopenfilename(filetypes=[("PDF File", "*.pdf")])

    if file:

        save = filedialog.asksaveasfilename(defaultextension=".docx")

        if save:

            def task():
                cv = Converter(file)
                cv.convert(save)
                cv.close()

            show_processing(task)

            messagebox.showinfo("Success", "PDF converted to Word")


# MERGE PDFs
def merge_pdfs():

    files = filedialog.askopenfilenames(filetypes=[("PDF Files", "*.pdf")])

    if files:

        save = filedialog.asksaveasfilename(defaultextension=".pdf")

        if save:

            def task():
                merger = PdfMerger()
                for pdf in files:
                    merger.append(pdf)
                merger.write(save)
                merger.close()

            show_processing(task)

            messagebox.showinfo("Success", "PDF files merged")


# MERGE WORD FILES
def merge_word():

    files = filedialog.askopenfilenames(filetypes=[("Word Files", "*.docx")])

    if files:

        save = filedialog.asksaveasfilename(defaultextension=".docx")

        if save:

            def task():
                merged = Document()

                for file in files:
                    doc = Document(file)
                    for para in doc.paragraphs:
                        merged.add_paragraph(para.text)

                merged.save(save)

            show_processing(task)

            messagebox.showinfo("Success", "Word files merged")


# MAIN WINDOW
root = tk.Tk()
root.title("Love to convert")
root.geometry("600x400")

title = tk.Label(root, text="Smart File Converter", font=("Arial", 20, "bold"))
title.pack(pady=20)

frame = tk.Frame(root)
frame.pack(pady=20)

tk.Button(frame, text="Word → PDF", width=20, height=2, command=word_to_pdf).grid(row=0, column=0, padx=20, pady=10)

tk.Button(frame, text="PDF → Word", width=20, height=2, command=pdf_to_word).grid(row=0, column=1, padx=20, pady=10)

tk.Button(frame, text="Merge PDFs", width=20, height=2, command=merge_pdfs).grid(row=1, column=0, padx=20, pady=10)

tk.Button(frame, text="Merge Word Files", width=20, height=2, command=merge_word).grid(row=1, column=1, padx=20, pady=10)

root.mainloop()